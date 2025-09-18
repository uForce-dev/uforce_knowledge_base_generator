import json
import logging
import re
import shutil
import time
from http import HTTPStatus
from pathlib import Path
from typing import Any

import requests
from docx import Document

from src.config import settings
from src.constants import TEAMLY_EXCLUDED_ARTICLE_IDS
from src.logging_config import setup_logging
from src.processors.base import BaseProcessor
from src.schemas import TeamlyArticle
from src.utils.gdrive_utils import (
    get_gdrive_service,
    upload_file_to_gdrive,
    delete_files_in_folder,
)


def clean_text(text: str) -> str:
    if not text:
        return ""

    emoji_pattern = re.compile(
        "["
        "\U0001f600-\U0001f64f"
        "\U0001f300-\U0001f5ff"
        "\U0001f680-\U0001f6ff"
        "\U0001f1e0-\U0001f1ff"
        "\U00002500-\U00002bef"
        "\U00002702-\U000027b0"
        "\U00002702-\U000027b0"
        "\U000024c2-\U0001f251"
        "\U0001f926-\U0001f937"
        "\U00010000-\U0010ffff"
        "\u2640-\u2642"
        "\u2600-\u2b55"
        "\u200d"
        "\u23cf"
        "\u23e9"
        "\u231a"
        "\ufe0f"
        "\u3030"
        "]+",
        flags=re.UNICODE,
    )
    text = emoji_pattern.sub(r"", text)
    text = re.sub(r"\s+", " ", text).strip()

    return text


def process_teamly_documents() -> None:
    """Entrypoint wrapper for class-based Teamly processing."""
    setup_logging()
    processor = TeamlyProcessor(logger=logging.getLogger(__name__))
    processor.run()


class TeamlyProcessor(BaseProcessor):
    TEAMLY_SLUG = settings.teamly_api_slug

    def __init__(
        self, logger: logging.Logger | None = None, use_cached_local_files: bool = False
    ) -> None:
        super().__init__(logger)
        # Token storage paths in secrets dir
        self._access_token_path: Path = settings.secrets_dir / "teamly_access_token.txt"
        self._refresh_token_path: Path = (
            settings.secrets_dir / "teamly_refresh_token.txt"
        )

        # Read tokens strictly from files; do not rely on .env values
        access_from_file = self._read_token_from_file(self._access_token_path)
        refresh_from_file = self._read_token_from_file(self._refresh_token_path)
        self._access_token = access_from_file
        self._refresh_token = refresh_from_file
        if not self._access_token or not self._refresh_token:
            self.logger.error(
                "Teamly tokens are missing. Please create and fill the following files:"
            )
            self.logger.error(
                f"  Access token file: {self._access_token_path} (put the raw access token string)"
            )
            self.logger.error(
                f"  Refresh token file: {self._refresh_token_path} (put the raw refresh token string)"
            )
        self._tokens_ready = bool(self._access_token and self._refresh_token)
        self._client_id = settings.teamly_api_client_id
        self._client_secret = settings.teamly_api_client_secret
        self._excluded_article_ids = TEAMLY_EXCLUDED_ARTICLE_IDS
        self._use_cached_local_files = use_cached_local_files

        self.authorize_endpoint = (
            f"https://{self.TEAMLY_SLUG}.teamly.ru/api/v1/auth/integration/authorize"
        )
        self.refresh_token_endpoint = (
            f"https://{self.TEAMLY_SLUG}.teamly.ru/api/v1/auth/integration/refresh"
        )
        self.articles_endpoint = f"https://{self.TEAMLY_SLUG}.teamly.ru/api/v1/integrations/space/{settings.teamly_space_id}/tree"
        self.article_detail_endpoint = (
            f"https://{self.TEAMLY_SLUG}.teamly.ru/api/v1/wiki/ql/article"
        )

    @property
    def headers(self) -> dict[str, str]:
        return {
            "Content-Type": "application/json",
            "X-Account-Slug": self.TEAMLY_SLUG,
            "Authorization": f"Bearer {self._access_token}",
        }

    def _persist_env_value(self, key: str, value: str) -> None:
        env_path: Path = settings.env_file
        try:
            if not env_path.exists():
                env_path.write_text(f"{key}={value}\n", encoding="utf-8")
                return
            content = env_path.read_text(encoding="utf-8")
            lines = content.splitlines()
            updated = False
            for idx, line in enumerate(lines):
                if line.strip().startswith(f"{key}="):
                    lines[idx] = f"{key}={value}"
                    updated = True
                    break
            if not updated:
                lines.append(f"{key}={value}")
            env_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        except Exception as exc:
            self.logger.warning(f"Failed to persist {key} to .env: {exc}")

    def _read_token_from_file(self, path: Path) -> str | None:
        try:
            if path.exists():
                value = path.read_text(encoding="utf-8").strip()
                return value or None
        except Exception as exc:
            self.logger.warning(f"Failed reading token file {path}: {exc}")
        return None

    def _write_token_to_file(self, path: Path, value: str) -> None:
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(value.strip(), encoding="utf-8")
        except Exception as exc:
            self.logger.warning(f"Failed writing token file {path}: {exc}")

    def _update_tokens_from_response(self, data: dict[str, Any]) -> bool:
        access = (
            data.get("access_token") or data.get("accessToken") or data.get("token")
        )
        refresh = data.get("refresh_token") or data.get("refreshToken")
        if not access:
            return False

        self._access_token = access
        if refresh:
            self._refresh_token = refresh

        # Persist to token files (preferred source)
        self._write_token_to_file(self._access_token_path, self._access_token)
        if refresh:
            self._write_token_to_file(self._refresh_token_path, self._refresh_token)

        self._persist_env_value("TEAMLY_API_ACCESS_TOKEN", self._access_token)
        if refresh:
            self._persist_env_value("TEAMLY_API_REFRESH_TOKEN", self._refresh_token)
        return True

    def refresh_token(self) -> dict[str, Any] | None:
        headers = {
            "Content-Type": "application/json",
            "X-Account-Slug": self.TEAMLY_SLUG,
        }
        response = requests.post(
            url=self.refresh_token_endpoint,
            json={
                "client_id": self._client_id,
                "client_secret": self._client_secret,
                "refresh_token": self._refresh_token,
            },
            headers=headers,
            timeout=30,
        )
        try:
            response.raise_for_status()
        except requests.HTTPError as exc:
            self.logger.error(f"Refresh token failed: {exc} | Body: {response.text}")
            return None
        response_json: dict[str, Any] = response.json()
        if not self._update_tokens_from_response(response_json):
            self.logger.error("Refresh token response missing expected fields.")
            return None
        self.logger.info("Teamly tokens refreshed and persisted.")
        return response_json

    def _request(self, method: str, url: str, **kwargs: Any) -> requests.Response:
        headers = kwargs.pop("headers", {}) or {}
        merged_headers = {**self.headers, **headers}
        timeout = kwargs.pop("timeout", 30)

        response = requests.request(
            method=method, url=url, headers=merged_headers, timeout=timeout, **kwargs
        )
        if response.status_code in (HTTPStatus.UNAUTHORIZED, HTTPStatus.FORBIDDEN):
            self.logger.info(
                "Access token likely expired. Attempting to refresh and retry..."
            )
            if self.refresh_token() is not None:
                merged_headers = {**self.headers, **headers}
                response = requests.request(
                    method=method,
                    url=url,
                    headers=merged_headers,
                    timeout=timeout,
                    **kwargs,
                )
        return response

    def list_articles_page(
        self, page: int = 1
    ) -> tuple[list[TeamlyArticle], dict[str, Any]]:
        time.sleep(0.2)
        self.logger.info(
            f"Fetching Teamly articles page={page} for space {settings.teamly_space_id}"
        )
        response = self._request(
            "GET",
            self.articles_endpoint,
            params={"page": page},
        )
        try:
            response.raise_for_status()
        except requests.HTTPError as exc:
            self.logger.error(
                f"Failed to fetch articles page {page}: {exc} | Body: {response.text}"
            )
            raise
        data = response.json() or {}
        items = data.get("items") or []
        pagination = data.get("pagination") or {}
        excluded = set(self._excluded_article_ids)
        if excluded:
            before = len(items)
            items = [it for it in items if str(it.get("id")) not in excluded]
            removed = before - len(items)
            if removed:
                self.logger.info(f"Excluded {removed} items on page {page}")
        parsed = [TeamlyArticle.model_validate(obj) for obj in items]
        self.logger.info(
            f"Fetched {len(parsed)} items on page {pagination.get('currentPage', page)} / {pagination.get('lastPage', '?')}"
        )
        return parsed, pagination

    def list_all_articles(self) -> list[TeamlyArticle]:
        page = 1
        all_items: list[TeamlyArticle] = []
        last_page: int | None = None
        while last_page is None or page <= last_page:
            items, pagination = self.list_articles_page(page=page)
            all_items.extend(items)
            current = int(pagination.get("currentPage", page) or page)
            last_page = int(pagination.get("lastPage", current) or current)
            self.logger.info(
                f"Accumulated articles: {len(all_items)} after page {current} of {last_page}"
            )
            if page >= last_page:
                break
            page += 1
        self.logger.info(f"Total articles collected: {len(all_items)}")
        return all_items

    def _extract_text_from_editor_content(self, content_field: Any) -> str:
        try:
            if isinstance(content_field, str):
                obj = json.loads(content_field)
            elif isinstance(content_field, dict):
                obj = content_field
            else:
                return ""

            parts: list[str] = []

            def visit(node: Any) -> None:
                if isinstance(node, dict):
                    text_val = node.get("text")
                    if isinstance(text_val, str):
                        parts.append(text_val)
                    for key in ("content", "children", "items", "paragraphs"):
                        child = node.get(key)
                        if isinstance(child, list):
                            for c in child:
                                visit(c)
                elif isinstance(node, list):
                    for item in node:
                        visit(item)

            visit(obj)
            combined = "\n".join(s for s in (p.strip() for p in parts) if s)
            return clean_text(combined)
        except Exception as exc:
            self.logger.warning(f"Failed to parse editor content: {exc}")
            return ""

    def get_article_details(self, article_id: str) -> dict[str, Any]:
        time.sleep(0.2)
        payload = {
            "query": {
                "__filter": {"id": article_id},
                "id": True,
                "title": True,
                "latestProperties": {"title": {"text": True}},
                "editorContentObject": {"content": True},
                "breadcrumbs": True,
                "relatedParentId": True,
            }
        }
        self.logger.info(f"Fetching Teamly article details for id={article_id}")
        try:
            response = self._request("POST", self.article_detail_endpoint, json=payload)
            response.raise_for_status()
        except requests.exceptions.SSLError as exc:
            self.logger.warning(
                f"SSL error while fetching article id={article_id}. Skipping. Error: {exc}"
            )
            return {}
        except requests.exceptions.RequestException as exc:
            self.logger.warning(
                f"Network error while fetching article id={article_id}. Skipping. Error: {exc}"
            )
            return {}
        except requests.HTTPError as exc:
            if response.status_code == HTTPStatus.NOT_FOUND:
                self.logger.warning(
                    f"Article not found (404) for id={article_id}. Skipping. Body: {response.text}"
                )
                return {}
            self.logger.error(
                f"Failed to fetch article {article_id}: {exc} | Body: {response.text}"
            )
            raise
        data = response.json() or {}
        self.logger.info(f"Fetched details for id={article_id}")
        return data

    def get_article_clean_text(self, article_id: str) -> str:
        data = self.get_article_details(article_id)
        editor_obj = (data or {}).get("editorContentObject") or {}
        content_field = editor_obj.get("content")
        text = self._extract_text_from_editor_content(content_field)
        if not text:
            latest = (data or {}).get("latestProperties") or {}
            title = ((latest.get("title") or {}).get("text")) or data.get("title")
            text = title or ""
        cleaned = clean_text(text)
        self.logger.info(
            f"Produced cleaned text for id={article_id}, length={len(cleaned)} chars"
        )
        return cleaned

    def get_article_clean_text_from_data(self, data: dict[str, Any]) -> str:
        editor_obj = (data or {}).get("editorContentObject") or {}
        content_field = editor_obj.get("content")
        text = self._extract_text_from_editor_content(content_field)
        if not text:
            latest = (data or {}).get("latestProperties") or {}
            title = ((latest.get("title") or {}).get("text")) or data.get("title")
            text = title or ""
        return clean_text(text)

    def _title_from_details(self, data: dict[str, Any]) -> str:
        latest = (data or {}).get("latestProperties") or {}
        title = ((latest.get("title") or {}).get("text")) or data.get("title") or ""
        return title

    def _top_level_id_from_details(self, article_id: str, data: dict[str, Any]) -> str:
        breadcrumbs = (data or {}).get("breadcrumbs") or []
        if isinstance(breadcrumbs, list) and breadcrumbs:
            top = breadcrumbs[0] or {}
            src = top.get("sourceId")
            if src:
                return str(src)
        parent = data.get("relatedParentId")
        if parent:
            return str(parent)
        return str(data.get("id") or article_id)

    def _second_level_id_from_details(self, data: dict[str, Any]) -> str | None:
        breadcrumbs = (data or {}).get("breadcrumbs") or []
        if isinstance(breadcrumbs, list) and len(breadcrumbs) >= 2:
            second = breadcrumbs[1] or {}
            src = second.get("sourceId")
            if src:
                return str(src)
        return None

    def _ancestor_ids_from_details(self, data: dict[str, Any]) -> list[str]:
        """Return all ancestor IDs from details payload (based on breadcrumbs and relatedParentId)."""
        ancestors: list[str] = []
        breadcrumbs = (data or {}).get("breadcrumbs") or []
        if isinstance(breadcrumbs, list):
            for node in breadcrumbs:
                if isinstance(node, dict):
                    src = node.get("sourceId")
                    if src:
                        ancestors.append(str(src))
        parent = (data or {}).get("relatedParentId")
        if parent:
            pid = str(parent)
            if pid not in ancestors:
                ancestors.append(pid)
        return ancestors

    def _is_excluded_or_descendant(self, article_id: str, data: dict[str, Any]) -> bool:
        excluded = set(self._excluded_article_ids)
        if article_id in excluded:
            return True
        for anc_id in self._ancestor_ids_from_details(data):
            if anc_id in excluded:
                return True
        return False

    def run(self) -> None:
        self.logger.info("Starting Teamly API processing...")
        if not getattr(self, "_tokens_ready", False):
            self.logger.error(
                "Aborting Teamly processing: token files are missing or empty. See errors above for instructions."
            )
            return
        service = get_gdrive_service()
        if not service:
            self.logger.error("Could not get Google Drive service. Aborting.")
            return
        processed_folder_id = settings.google_drive_teamly_processed_dir_id
        temp_dir = settings.teamly_temp_dir
        temp_dir.mkdir(parents=True, exist_ok=True)
        if not self._use_cached_local_files:
            removed = 0
            for item in temp_dir.iterdir():
                try:
                    if item.is_file() or item.is_symlink():
                        item.unlink()
                        removed += 1
                    elif item.is_dir():
                        shutil.rmtree(item)
                        removed += 1
                except Exception as e:
                    self.logger.warning(f"Failed to remove {item}: {e}")
            self.logger.info(
                f"Cleared local Teamly temp dir {temp_dir}, removed {removed} items"
            )
        self.logger.info(f"Clearing processed folder ID: {processed_folder_id}...")
        delete_files_in_folder(service, processed_folder_id)
        self.logger.info("Processed folder cleared.")

        articles: list[TeamlyArticle] = []
        if not self._use_cached_local_files:
            articles = self.list_all_articles()
            if not articles:
                self.logger.info("No articles found to process.")
                return

        details_cache: dict[str, dict[str, Any]] = {}
        cleaned_text_cache: dict[str, str] = {}
        groups: dict[str, list[str]] = {}
        group_titles: dict[str, str] = {}
        combined_txt_paths: list[Path] = []

        if not self._use_cached_local_files:
            total = len(articles)
            for idx, art in enumerate(articles, start=1):
                self.logger.info(
                    f"Details progress {idx}/{total} ({(idx / total) * 100:.1f}%) id={art.id}"
                )
                data = self.get_article_details(art.id)
                if not data:
                    continue
                # Skip any article that is excluded explicitly or is a descendant of an excluded node
                try:
                    if self._is_excluded_or_descendant(art.id, data):
                        self.logger.info(
                            f"Skipping id={art.id} because it or one of its ancestors is excluded"
                        )
                        continue
                except Exception:
                    # Fail-open: if we cannot determine ancestry, proceed
                    pass
                details_cache[art.id] = data
                second_id = self._second_level_id_from_details(data)
                if second_id:
                    groups.setdefault(second_id, []).append(art.id)
                    if second_id not in group_titles:
                        # fetch title for the second-level node
                        if second_id in details_cache:
                            group_titles[second_id] = self._title_from_details(
                                details_cache[second_id]
                            )
                        else:
                            try:
                                top2 = self.get_article_details(second_id)
                                if top2:
                                    details_cache[second_id] = top2
                                    group_titles[second_id] = self._title_from_details(
                                        top2
                                    )
                            except Exception:
                                group_titles[second_id] = ""
                # cache cleaned text
                cleaned_text_cache[art.id] = self.get_article_clean_text_from_data(data)

        # If using cached local files, rebuild grouping from local .txt files
        if self._use_cached_local_files:
            groups.clear()
            group_titles.clear()
            for path in sorted(temp_dir.glob("*.txt")):
                try:
                    text = path.read_text(encoding="utf-8")
                except Exception as e:
                    self.logger.warning(f"Failed reading {path}: {e}")
                    continue
                folder = None
                title = path.stem
                if text.startswith("---\n"):
                    end = text.find("\n---\n")
                    header = text[4:end] if end != -1 else ""
                    for line in header.splitlines():
                        if line.startswith("folder:"):
                            folder = line.split(":", 1)[1].strip()
                            break
                if not folder:
                    folder = "Teamly"
                group_titles.setdefault(folder, folder)
                # Body without metadata and first heading
                body = text
                if text.startswith("---\n") and end != -1:
                    body = text[end + len("\n---\n") :]
                lines = body.splitlines()
                if lines and lines[0].startswith("# "):
                    body = "\n".join(lines[1:])
                groups.setdefault(folder, []).append((title, body))

        if not self._use_cached_local_files:
            # Build combined TXT per second-level group; exclude the group node itself
            for second_id, article_ids in groups.items():
                child_ids = [aid for aid in article_ids if aid != second_id]
                if not child_ids:
                    continue
                combined_chunks: list[tuple[str, str]] = []
                for aid in child_ids:
                    data = details_cache.get(aid)
                    if data is None:
                        data = self.get_article_details(aid)
                        if not data:
                            continue
                        details_cache[aid] = data
                    text = cleaned_text_cache.get(aid)
                    if text is None:
                        text = self.get_article_clean_text_from_data(data)
                        cleaned_text_cache[aid] = text
                    if not text:
                        continue
                    title = self._title_from_details(details_cache.get(aid, {})) or aid
                    combined_chunks.append((title, text))

                if not combined_chunks:
                    continue
                folder_name = group_titles.get(second_id) or second_id
                safe_folder_name = re.sub(r"[^\w\-_. ]+", "_", folder_name).replace(
                    " ", "_"
                )
                combined_txt = temp_dir / f"teamly__{safe_folder_name}.txt"
                try:
                    with open(combined_txt, "w", encoding="utf-8") as f:
                        f.write("---\n")
                        f.write("source: Teamly\n")
                        f.write(f"folder: {folder_name}\n")
                        f.write("tz: Europe/Moscow\n")
                        f.write("body_format: kv-blocks\n")
                        f.write("---\n\n")
                        for title, content_str in combined_chunks:
                            f.write(f"# {title}\n")
                            f.write(content_str.rstrip("\n") + "\n\n")
                    combined_txt_paths.append(combined_txt)
                    self.logger.info(
                        f"Generated combined TXT for '{folder_name}': {combined_txt}"
                    )
                except Exception as e:
                    self.logger.error(
                        f"Error writing combined TXT {combined_txt} for '{folder_name}': {e}"
                    )
        else:
            # Build combined docs purely from cached local files
            for folder_name, items in groups.items():
                # Exclude item whose title equals folder_name (top-level itself)
                combined_chunks = [
                    (title, body) for title, body in items if title != folder_name
                ]
                if not combined_chunks:
                    continue
                safe_folder_name = re.sub(r"[^\w\-_. ]+", "_", folder_name).replace(
                    " ", "_"
                )
                combined_txt = temp_dir / f"teamly__{safe_folder_name}.txt"
                try:
                    with open(combined_txt, "w", encoding="utf-8") as f:
                        f.write("---\n")
                        f.write("source: Teamly\n")
                        f.write(f"folder: {folder_name}\n")
                        f.write("tz: Europe/Moscow\n")
                        f.write("body_format: kv-blocks\n")
                        f.write("---\n\n")
                        for title, content_str in combined_chunks:
                            f.write(f"# {title}\n")
                            f.write(content_str.rstrip("\n") + "\n\n")
                    combined_txt_paths.append(combined_txt)
                    self.logger.info(
                        f"Generated combined TXT for '{folder_name}': {combined_txt}"
                    )
                except Exception as e:
                    self.logger.error(
                        f"Error writing combined TXT {combined_txt} for '{folder_name}': {e}"
                    )

        # Convert generated TXT files (or all in cached mode) to DOCX and upload
        txt_targets: list[Path]
        if self._use_cached_local_files:
            txt_targets = list(sorted(temp_dir.glob("*.txt")))
        else:
            txt_targets = combined_txt_paths
        for txt_path in txt_targets:
            try:
                folder_name = txt_path.stem
                docx_path = temp_dir / f"{folder_name}.docx"
                text = txt_path.read_text(encoding="utf-8")
                doc = Document()
                for line in text.splitlines():
                    if line.startswith("# "):
                        doc.add_heading(line[2:].strip(), level=1)
                    else:
                        doc.add_paragraph(line)
                doc.save(docx_path)
                self.logger.info(f"Converted TXT to DOCX: {docx_path}")
                self.logger.info(
                    f"Uploading {docx_path.stem} to Google Drive as a Google Doc..."
                )
                upload_file_to_gdrive(
                    service, docx_path, processed_folder_id, as_gdoc=True
                )
            except Exception as e:
                self.logger.error(
                    f"Failed converting/uploading {txt_path.name} to DOCX/GDoc: {e}"
                )


if __name__ == "__main__":
    process_teamly_documents()
